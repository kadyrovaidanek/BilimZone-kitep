from io import BytesIO
import os
import traceback

from django.http import HttpResponse

from docx import Document
from openpyxl import Workbook

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

from reports.services import get_admin_report_data, get_owner_report_data


PDF_FONT_NAME = "Helvetica"


def money(value):
    return f"{float(value or 0):,.2f} сом".replace(",", " ")


def make_response(content, filename, content_type):
    response = HttpResponse(content, content_type=content_type)
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


def get_report_data(role, request, user_id=None):
    if role == "admin":
        return get_admin_report_data(request)

    return get_owner_report_data(request, user_id)


def register_pdf_font():
    global PDF_FONT_NAME

    font_paths = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/Library/Fonts/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
    ]

    for font_path in font_paths:
        if os.path.exists(font_path):
            pdfmetrics.registerFont(TTFont("BilimZoneFont", font_path))
            PDF_FONT_NAME = "BilimZoneFont"
            return PDF_FONT_NAME

    return PDF_FONT_NAME


def paragraph_cell(value, style):
    return Paragraph(str(value), style)


def build_summary_rows(data):
    summary = data.get("summary", {})

    if data.get("role") == "admin":
        return [
            ["Общая сумма продаж", money(summary.get("total_sales"))],
            ["Доход платформы", money(summary.get("platform_income"))],
            ["Доход авторов/организаций", money(summary.get("owners_income"))],
            ["Количество покупок", str(summary.get("purchases_count", 0))],
            ["Комиссия платформы", f'{data.get("commission_percent", "0")}%'],
        ]

    return [
        ["Мой доход", money(summary.get("owner_income"))],
        ["Общая сумма продаж", money(summary.get("total_sales"))],
        ["Количество покупок", str(summary.get("purchases_count", 0))],
        ["Лучший материал", summary.get("top_material") or "Нет данных"],
        ["Комиссия платформы", f'{data.get("commission_percent", "0")}%'],
    ]


def build_publication_rows(data):
    tables = data.get("tables", {})

    if data.get("role") == "admin":
        rows = [
            [
                "Материал",
                "Категория",
                "Владелец",
                "Покупки",
                "Общая сумма",
                "Доход платформы",
                "Доход владельца",
            ]
        ]

        for item in tables.get("top_publications", []):
            rows.append(
                [
                    item.get("title", ""),
                    item.get("category") or "Без категории",
                    item.get("owner", ""),
                    item.get("purchases_count", 0),
                    money(item.get("total_sales")),
                    money(item.get("platform_income")),
                    money(item.get("owners_income")),
                ]
            )

        return rows

    rows = [
        [
            "Материал",
            "Категория",
            "Покупки",
            "Общая сумма",
            "Мой доход",
            "Комиссия платформы",
        ]
    ]

    for item in tables.get("publications", []):
        rows.append(
            [
                item.get("title", ""),
                item.get("category") or "Без категории",
                item.get("purchases_count", 0),
                money(item.get("total_sales")),
                money(item.get("owner_income")),
                money(item.get("platform_commission")),
            ]
        )

    return rows


def build_owner_rows(data):
    if data.get("role") != "admin":
        return []

    rows = [
        [
            "Автор/организация",
            "Покупки",
            "Общая сумма",
            "Доход платформы",
            "Доход владельца",
        ]
    ]

    for item in data.get("tables", {}).get("top_owners", []):
        rows.append(
            [
                item.get("owner", ""),
                item.get("purchases_count", 0),
                money(item.get("total_sales")),
                money(item.get("platform_income")),
                money(item.get("owner_income")),
            ]
        )

    return rows


def export_report_excel(role, request, user_id=None):
    data = get_report_data(role, request, user_id)

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Report"

    title = "Отчёт администратора" if role == "admin" else "Отчёт автора / организации"

    sheet.append([title])
    sheet.append([])

    sheet.append(["Показатель", "Значение"])
    for row in build_summary_rows(data):
        sheet.append(row)

    sheet.append([])
    sheet.append(["Материалы"])
    for row in build_publication_rows(data):
        sheet.append(row)

    owner_rows = build_owner_rows(data)
    if owner_rows:
        sheet.append([])
        sheet.append(["Авторы / организации"])
        for row in owner_rows:
            sheet.append(row)

    for column_cells in sheet.columns:
        max_length = 0
        column_letter = column_cells[0].column_letter

        for cell in column_cells:
            value = str(cell.value or "")
            max_length = max(max_length, len(value))

        sheet.column_dimensions[column_letter].width = min(max_length + 4, 45)

    buffer = BytesIO()
    workbook.save(buffer)
    buffer.seek(0)

    filename = "admin_report.xlsx" if role == "admin" else "owner_report.xlsx"

    return make_response(
        buffer.getvalue(),
        filename,
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


def export_report_docx(role, request, user_id=None):
    data = get_report_data(role, request, user_id)

    document = Document()

    title = "Отчёт администратора" if role == "admin" else "Отчёт автора / организации"

    document.add_heading(title, 0)
    document.add_paragraph(f'Комиссия платформы: {data.get("commission_percent", "0")}%')

    document.add_heading("Основные показатели", level=1)

    summary_table = document.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"
    summary_table.rows[0].cells[0].text = "Показатель"
    summary_table.rows[0].cells[1].text = "Значение"

    for label, value in build_summary_rows(data):
        row = summary_table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value)

    document.add_heading("Материалы", level=1)

    publication_rows = build_publication_rows(data)

    if not publication_rows:
        publication_rows = [["Материал"], ["Нет данных"]]

    publication_table = document.add_table(rows=1, cols=len(publication_rows[0]))
    publication_table.style = "Table Grid"

    for index, value in enumerate(publication_rows[0]):
        publication_table.rows[0].cells[index].text = str(value)

    for source_row in publication_rows[1:]:
        row = publication_table.add_row().cells

        for index, value in enumerate(source_row):
            row[index].text = str(value)

    owner_rows = build_owner_rows(data)

    if owner_rows:
        document.add_heading("Авторы / организации", level=1)

        owner_table = document.add_table(rows=1, cols=len(owner_rows[0]))
        owner_table.style = "Table Grid"

        for index, value in enumerate(owner_rows[0]):
            owner_table.rows[0].cells[index].text = str(value)

        for source_row in owner_rows[1:]:
            row = owner_table.add_row().cells

            for index, value in enumerate(source_row):
                row[index].text = str(value)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    filename = "admin_report.docx" if role == "admin" else "owner_report.docx"

    return make_response(
        buffer.getvalue(),
        filename,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def export_report_pdf(role, request, user_id=None):
    try:
        data = get_report_data(role, request, user_id)
        font_name = register_pdf_font()

        buffer = BytesIO()

        pdf = SimpleDocTemplate(
            buffer,
            pagesize=landscape(A4),
            rightMargin=24,
            leftMargin=24,
            topMargin=24,
            bottomMargin=24,
        )

        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            "BilimZoneTitle",
            parent=styles["Title"],
            fontName=font_name,
            fontSize=20,
            leading=24,
            alignment=1,
        )

        heading_style = ParagraphStyle(
            "BilimZoneHeading",
            parent=styles["Heading2"],
            fontName=font_name,
            fontSize=14,
            leading=18,
            spaceAfter=8,
        )

        table_style = ParagraphStyle(
            "BilimZoneTable",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=8,
            leading=10,
        )

        small_table_style = ParagraphStyle(
            "BilimZoneSmallTable",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=7,
            leading=9,
        )

        elements = []

        title = "Отчёт администратора" if role == "admin" else "Отчёт автора / организации"

        elements.append(Paragraph(title, title_style))
        elements.append(Spacer(1, 16))

        elements.append(Paragraph("Основные показатели", heading_style))

        summary_rows = [
            [
                paragraph_cell("Показатель", table_style),
                paragraph_cell("Значение", table_style),
            ]
        ]

        for label, value in build_summary_rows(data):
            summary_rows.append(
                [
                    paragraph_cell(label, table_style),
                    paragraph_cell(value, table_style),
                ]
            )

        summary_table = Table(summary_rows, repeatRows=1, colWidths=[220, 260])
        summary_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("FONTNAME", (0, 0), (-1, -1), font_name),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]
            )
        )

        elements.append(summary_table)
        elements.append(Spacer(1, 18))

        elements.append(Paragraph("Материалы", heading_style))

        raw_publication_rows = build_publication_rows(data)

        publication_rows = []
        for row in raw_publication_rows:
            publication_rows.append(
                [paragraph_cell(value, small_table_style) for value in row]
            )

        if not publication_rows:
            publication_rows = [[paragraph_cell("Нет данных", small_table_style)]]

        publication_table = Table(publication_rows, repeatRows=1)
        publication_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("FONTNAME", (0, 0), (-1, -1), font_name),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )

        elements.append(publication_table)

        owner_rows = build_owner_rows(data)

        if owner_rows:
            elements.append(Spacer(1, 18))
            elements.append(Paragraph("Авторы / организации", heading_style))

            prepared_owner_rows = []
            for row in owner_rows:
                prepared_owner_rows.append(
                    [paragraph_cell(value, small_table_style) for value in row]
                )

            owner_table = Table(prepared_owner_rows, repeatRows=1)
            owner_table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("FONTNAME", (0, 0), (-1, -1), font_name),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ]
                )
            )

            elements.append(owner_table)

        pdf.build(elements)
        buffer.seek(0)

        filename = "admin_report.pdf" if role == "admin" else "owner_report.pdf"

        return make_response(buffer.getvalue(), filename, "application/pdf")

    except Exception as error:
        print("REPORT PDF EXPORT ERROR:", error)
        print(traceback.format_exc())

        return HttpResponse(
            "PDF export error. Please use Excel or DOCX export.",
            status=500,
            content_type="text/plain; charset=utf-8",
        )


def export_report(role, export_format, request, user_id=None):
    if export_format == "excel":
        return export_report_excel(role, request, user_id)

    if export_format == "docx":
        return export_report_docx(role, request, user_id)

    if export_format == "pdf":
        return export_report_pdf(role, request, user_id)

    return HttpResponse(
        "Unsupported export format",
        status=400,
        content_type="text/plain; charset=utf-8",
    )